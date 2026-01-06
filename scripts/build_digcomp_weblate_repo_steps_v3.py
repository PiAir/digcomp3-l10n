#!/usr/bin/env python3
"""
build_digcomp_weblate_repo_steps_v3.py

One script to:
  1) Create ./digcomp3-l10n/ repo structure + extract "normative" strings from XLSX into locale/*/en.csv (+ nl.csv copies)
  2) Generate a *starter* manifest.json (based on a known DigComp 3.0 structure; you will likely tweak headings once)
  3) Extract text blocks from DOCX into locale/texts/en.csv and merge/update locale/texts/nl.csv (preserving existing NL)

Key improvements vs v2
- Robust CSV writer: always quotes commas/newlines correctly (prevents broken rows in Weblate)
- Context column in every CSV row (helps translators & later DOCX rebuild)
- DOCX extraction walks the document body *in order* and includes table cell text (python-docx does NOT include table text in doc.paragraphs)
- Conservative text cleanup:
    * normalise repeated parentheses created by earlier extraction quirks
    * strip superfluous whitespace while preserving intentional line breaks

Folder layout
- This script can live in the top folder next to the source XLSX/DOCX.
- It writes the repo into ./digcomp3-l10n/ and CSVs into ./digcomp3-l10n/locale/<component>/

Usage
  python build_digcomp_weblate_repo_steps_v3.py step1 --xlsx "DigComp 3.0 Data Supplement 24 Nov 2025.xlsx"
  python build_digcomp_weblate_repo_steps_v3.py step2 --jsonld "DigComp 3.0 Data Supplement 24 Nov 2025.jsonld"
  python build_digcomp_weblate_repo_steps_v3.py step3 --docx "DigComp 3.0 EDITABLE 16 Dec 2025.docx"

Step meanings
- step1: create repo structure + extract xlsx -> core-framework/levels/statements/outcomes/glossary (en.csv + nl.csv stub)
- step2: consistency check against JSON-LD (reports missing/extra statement/outcome IDs, etc.)
- step3: generate manifest.json if missing and extract DOCX text blocks per manifest into locale/texts/

Note
- step2 is read-only.
- step3 relies on manifest.json; if you change headings in the DOCX, update manifest anchors.

"""

from __future__ import annotations

import argparse
import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, Iterator, List, Optional, Tuple

from openpyxl import load_workbook
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


# ----------------------------
# Repo skeleton
# ----------------------------

COMPONENTS = [
    "acronyms",
    "core-framework",
    "glossary",
    "levels",
    "outcomes",
    "statements",
    "texts",
]


def ensure_repo(repo_root: Path) -> None:
    (repo_root / "digcomp3-l10n" / "locale").mkdir(parents=True, exist_ok=True)
    for comp in COMPONENTS:
        (repo_root / "digcomp3-l10n" / "locale" / comp).mkdir(parents=True, exist_ok=True)


# ----------------------------
# CSV I/O
# ----------------------------

CSV_FIELDS = ["location", "source", "target", "context"]


def write_csv(path: Path, rows: List[Dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=CSV_FIELDS, quoting=csv.QUOTE_MINIMAL)
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k, "") for k in CSV_FIELDS})


def read_csv(path: Path) -> Dict[str, Dict[str, str]]:
    out: Dict[str, Dict[str, str]] = {}
    if not path.exists():
        return out
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        rdr = csv.DictReader(f)
        for row in rdr:
            loc = (row.get("location") or "").strip()
            if not loc:
                continue
            out[loc] = {
                "location": loc,
                "source": row.get("source","") or "",
                "target": row.get("target","") or "",
                "context": row.get("context","") or "",
            }
    return out


def upsert_locale_csv(en_path: Path, nl_path: Path, new_rows: List[Dict[str, str]]) -> None:
    """
    Write en.csv fresh, and merge into nl.csv while preserving existing NL targets.
    """
    # write EN
    write_csv(en_path, new_rows)

    # merge NL
    existing = read_csv(nl_path)
    merged: List[Dict[str, str]] = []
    preserved = 0
    for r in new_rows:
        loc = r["location"]
        tgt = ""
        if loc in existing and (existing[loc].get("target") or "").strip():
            tgt = existing[loc]["target"]
            preserved += 1
        merged.append({
            "location": loc,
            "source": r["source"],
            "target": tgt,
            "context": r.get("context",""),
        })
    write_csv(nl_path, merged)
    print(f"[OK] Wrote: {en_path} ({len(new_rows)} rows)")
    print(f"[OK] Wrote: {nl_path} ({len(merged)} rows) – preserved {preserved} existing targets")


# ----------------------------
# XLSX extraction (step1)
# ----------------------------

def norm_num(x) -> str:
    if x is None:
        return ""
    if isinstance(x, int):
        return str(x)
    if isinstance(x, float):
        if x.is_integer():
            return str(int(x))
        return f"{x:.10f}".rstrip("0").rstrip(".")
    s = str(x).strip()
    if re.match(r"^\d+\.0$", s):
        return s[:-2]
    return s


def slugify_term(term: str) -> str:
    s = term.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s


def step1_extract_xlsx(xlsx_path: Path, repo_root: Path) -> None:
    ensure_repo(repo_root)
    wb = load_workbook(xlsx_path, data_only=True)

    locale = repo_root / "digcomp3-l10n" / "locale"

    # --- core-framework from sheet 1
    ws = wb["1 Competence Areas&Competences"]
    rows: List[Dict[str, str]] = []
    seen = set()

    for r in range(2, ws.max_row + 1):
        area_no = norm_num(ws.cell(r, 1).value)
        area_name = (ws.cell(r, 2).value or "").strip()
        area_desc = (ws.cell(r, 3).value or "").strip()
        comp_no = norm_num(ws.cell(r, 4).value)
        comp_name = (ws.cell(r, 5).value or "").strip()
        comp_desc = (ws.cell(r, 6).value or "").strip()

        def add(loc, src, ctx):
            if not loc or loc in seen:
                return
            seen.add(loc)
            rows.append({"location": loc, "source": src, "target": "", "context": ctx})

        add(f"digcomp.area.{area_no}.label", area_name, f"Competence area {area_no} – label")
        add(f"digcomp.area.{area_no}.description", area_desc, f"Competence area {area_no} – description")
        add(f"digcomp.competence.{comp_no}.label", comp_name, f"Competence {comp_no} – label")
        add(f"digcomp.competence.{comp_no}.description", comp_desc, f"Competence {comp_no} – description")

    upsert_locale_csv(locale/"core-framework"/"en.csv", locale/"core-framework"/"nl.csv", rows)

    # --- levels from sheet 2
    ws = wb["2 Proficiency Levels"]
    rows = []
    for r in range(2, ws.max_row + 1):
        eight = ws.cell(r, 4).value
        if eight is None:
            continue
        eight_s = norm_num(eight)
        four_name = (ws.cell(r, 1).value or "").strip()
        four_desc = (ws.cell(r, 2).value or "").strip()
        applies = (ws.cell(r, 3).value or "").strip()
        eight_desc = (ws.cell(r, 6).value or "").strip()
        rows.extend([
            {"location": f"digcomp.level.{eight_s}.four_level_name", "source": four_name, "target": "", "context": f"Proficiency level {eight_s} – 4-level name"},
            {"location": f"digcomp.level.{eight_s}.four_level_description", "source": four_desc, "target": "", "context": f"Proficiency level {eight_s} – 4-level description"},
            {"location": f"digcomp.level.{eight_s}.applies_to", "source": applies, "target": "", "context": f"Proficiency level {eight_s} – applies to"},
            {"location": f"digcomp.level.{eight_s}.eight_level_description", "source": eight_desc, "target": "", "context": f"Proficiency level {eight_s} – 8-level description"},
        ])
    upsert_locale_csv(locale/"levels"/"en.csv", locale/"levels"/"nl.csv", rows)

    # --- statements from sheet 3
    ws = wb["3 Competence Statements"]
    rows = []
    for r in range(2, ws.max_row + 1):
        sid = str(ws.cell(r, 7).value or "").strip()
        txt = (ws.cell(r, 8).value or "").strip()
        if not sid or not txt:
            continue
        rows.append({"location": f"digcomp.statement.{sid}", "source": txt, "target": "", "context": f"Competence statement {sid}"})
    upsert_locale_csv(locale/"statements"/"en.csv", locale/"statements"/"nl.csv", rows)

    # --- outcomes from sheet 4
    ws = wb["4 Learning Outcomes"]
    rows = []
    for r in range(2, ws.max_row + 1):
        oid = str(ws.cell(r, 5).value or "").strip()
        txt = (ws.cell(r, 6).value or "").strip()
        if not oid or not txt:
            continue
        rows.append({"location": f"digcomp.outcome.{oid}", "source": txt, "target": "", "context": f"Learning outcome {oid}"})
    upsert_locale_csv(locale/"outcomes"/"en.csv", locale/"outcomes"/"nl.csv", rows)

    # --- glossary from sheet 5
    ws = wb["5 Glossary"]
    rows = []
    for r in range(2, ws.max_row + 1):
        term = (ws.cell(r, 1).value or "").strip()
        expl = (ws.cell(r, 2).value or "").strip()
        if not term:
            continue
        slug = slugify_term(term)
        rows.append({"location": f"digcomp.glossary.{slug}.label", "source": term, "target": "", "context": "Glossary term"})
        rows.append({"location": f"digcomp.glossary.{slug}.definition", "source": expl, "target": "", "context": f"Glossary definition for {term}"})
    upsert_locale_csv(locale/"glossary"/"en.csv", locale/"glossary"/"nl.csv", rows)

    # Acronyms & texts are filled later (step3) – keep empty files so Weblate can create components consistently
    for comp in ("acronyms", "texts"):
        en = locale/comp/"en.csv"
        nl = locale/comp/"nl.csv"
        if not en.exists():
            write_csv(en, [])
        if not nl.exists():
            write_csv(nl, [])
    print("[OK] Step1 complete.")


# ----------------------------
# JSON-LD consistency check (step2)
# ----------------------------

def step2_check_jsonld(jsonld_path: Path, repo_root: Path) -> None:
    locale = repo_root / "digcomp3-l10n" / "locale"
    stm = read_csv(locale/"statements"/"en.csv")
    out = read_csv(locale/"outcomes"/"en.csv")
    # sets from CSV
    stmt_ids = {k.split(".",2)[2] for k in stm.keys() if k.startswith("digcomp.statement.")}
    out_ids = {k.split(".",2)[2] for k in out.keys() if k.startswith("digcomp.outcome.")}

    data = json.loads(jsonld_path.read_text(encoding="utf-8"))
    graph = data.get("@graph", [])
    jl_stmt = set()
    jl_out = set()
    for node in graph:
        t = node.get("@type")
        _id = str(node.get("@id",""))
        if t == "CompetenceStatement" and "/" in _id:
            jl_stmt.add(_id.split("/",1)[1])
        if t == "LearningOutcome" and "/" in _id:
            jl_out.add(_id.split("/",1)[1])

    missing_stmt = sorted(jl_stmt - stmt_ids)
    missing_out = sorted(jl_out - out_ids)
    extra_stmt = sorted(stmt_ids - jl_stmt)
    extra_out = sorted(out_ids - jl_out)

    print("JSON-LD consistency report")
    print(f"- Competence statements: JSON-LD={len(jl_stmt)} CSV={len(stmt_ids)} missing_in_CSV={len(missing_stmt)} extra_in_CSV={len(extra_stmt)}")
    if missing_stmt[:20]:
        print("  examples missing_in_CSV:", missing_stmt[:20])
    print(f"- Learning outcomes: JSON-LD={len(jl_out)} CSV={len(out_ids)} missing_in_CSV={len(missing_out)} extra_in_CSV={len(extra_out)}")
    if missing_out[:20]:
        print("  examples missing_in_CSV:", missing_out[:20])
    print("[OK] Step2 complete.")


# ----------------------------
# DOCX text extraction (step3)
# ----------------------------

def iter_block_items(doc: Document) -> Iterator[object]:
    """
    Yield paragraphs and tables in document order.
    Source: common python-docx pattern.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def clean_text(s: str) -> str:
    # normalize whitespace, but keep intentional newlines
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # collapse repeated parentheses artefacts like "(((" or ")))"
    s = re.sub(r"\({2,}", "(", s)
    s = re.sub(r"\){2,}", ")", s)
    # collapse spaces
    s = re.sub(r"[ \t]+", " ", s)
    # trim each line
    s = "\n".join([ln.strip() for ln in s.split("\n")])
    return s.strip()


def find_anchor_index(blocks: List[Tuple[str,str,str]], anchor: Dict) -> Optional[int]:
    """
    blocks: list of (kind, style, text)
    anchor: {"type": "heading_contains", "value": "..."}
    Return index of the anchor block.
    """
    atype = (anchor or {}).get("type")
    aval = (anchor or {}).get("value","").strip().lower()
    if not aval:
        return None
    for i, (kind, style, text) in enumerate(blocks):
        if kind != "p":
            continue
        if not style.lower().startswith("heading"):
            continue
        t = (text or "").lower()
        if atype == "heading_contains" and aval in t:
            return i
    return None


def step3_extract_texts(docx_path: Path, repo_root: Path, manifest_path: Optional[Path] = None) -> None:
    ensure_repo(repo_root)
    if manifest_path is None:
        manifest_path = repo_root / "manifest.json"
    if not manifest_path.exists():
        # write a conservative starter manifest matching the earlier agreed structure
        starter = {
            "version": 1,
            "notes": "Starter manifest generated by build_digcomp_weblate_repo_steps_v3.py; adjust anchors if needed.",
            "sections": [
                {"id":"front.colophon","title":"Colophon","action":"import","docx_anchor":{"type":"heading_contains","value":"colophon"},"import":{"key_prefix":"doc.front.colophon"}},
                {"id":"front.abstract","title":"Abstract","action":"import","docx_anchor":{"type":"heading_contains","value":"abstract"},"import":{"key_prefix":"doc.front.abstract"}},
                {"id":"front.foreword","title":"Foreword","action":"import","docx_anchor":{"type":"heading_contains","value":"foreword"},"import":{"key_prefix":"doc.front.foreword"}},
                {"id":"front.acknowledgements","title":"Acknowledgements","action":"import","docx_anchor":{"type":"heading_contains","value":"acknowledgements"},"import":{"key_prefix":"doc.front.acknowledgements"}},
                {"id":"front.executive_summary","title":"Executive summary","action":"import","docx_anchor":{"type":"heading_contains","value":"executive summary"},"import":{"key_prefix":"doc.front.executive_summary"}},
                {"id":"front.quick_guide","title":"Quick guide to DigComp 3.0","action":"import","docx_anchor":{"type":"heading_contains","value":"quick guide"},"import":{"key_prefix":"doc.front.quick_guide"}},

                {"id":"ch1.introduction","title":"1. INTRODUCTION","action":"import","docx_anchor":{"type":"heading_contains","value":"1. introduction"},"import":{"key_prefix":"doc.ch1"}},

                {"id":"ch2.framework_components","title":"2. DIGCOMP 3.0 FRAMEWORK COMPONENTS","action":"import_export","docx_anchor":{"type":"heading_contains","value":"framework components"},"import":{"key_prefix":"doc.ch2"},"export":[{"component":"levels"}]},

                {"id":"ch3.how_to_read","title":"3.1 How to read DigComp 3.0","action":"import","docx_anchor":{"type":"heading_contains","value":"how to read digcomp 3.0"},"import":{"key_prefix":"doc.ch3.how_to_read"}},
                {"id":"ch3.framework","title":"3. DIGCOMP 3.0 FRAMEWORK","action":"import_export","docx_anchor":{"type":"heading_contains","value":"digcomp 3.0 framework"},"import":{"key_prefix":"doc.ch3"},"export":[{"component":"core-framework"},{"component":"statements"}]},

                {"id":"ch4.concluding_remarks","title":"4. CONCLUDING REMARKS","action":"import","docx_anchor":{"type":"heading_contains","value":"concluding remarks"},"import":{"key_prefix":"doc.ch4"}},

                {"id":"list.acronyms","title":"LIST OF ACRONYMS","action":"export","docx_anchor":{"type":"heading_contains","value":"list of acronyms"},"export":[{"component":"acronyms"}]},
                {"id":"glossary","title":"GLOSSARY OF TERMS AND DEFINITIONS","action":"export","docx_anchor":{"type":"heading_contains","value":"glossary of terms"},"export":[{"component":"glossary"}]},

                {"id":"annex2.learning_outcomes","title":"Annex 2: DigComp 3.0 learning outcomes","action":"import_export","docx_anchor":{"type":"heading_contains","value":"annex 2"},"import":{"key_prefix":"doc.annex2"},"export":[{"component":"outcomes"}]},

                {"id":"annex3.phases","title":"Annex 3: Phases in the development of DigComp 3.0","action":"import","docx_anchor":{"type":"heading_contains","value":"annex 3"},"import":{"key_prefix":"doc.annex3"}},
            ],
        }
        manifest_path.write_text(json.dumps(starter, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[OK] Wrote starter manifest: {manifest_path}")

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    sections = manifest.get("sections", [])

    doc = Document(docx_path)

    # Build ordered blocks with style and text
    blocks: List[Tuple[str,str,str]] = []
    for blk in iter_block_items(doc):
        if isinstance(blk, Paragraph):
            blocks.append(("p", blk.style.name if blk.style else "", clean_text(blk.text or "")))
        elif isinstance(blk, Table):
            # Represent a table block; we keep text in cells later when assigning
            blocks.append(("tbl", "Table", ""))

    # Resolve section boundaries by anchors
    anchors: Dict[str, int] = {}
    for sec in sections:
        idx = find_anchor_index(blocks, sec.get("docx_anchor") or {})
        if idx is not None:
            anchors[sec["id"]] = idx

    # Warn about missing anchors
    missing = [sec["id"] for sec in sections if sec.get("docx_anchor") and sec["id"] not in anchors]
    if missing:
        print("[WARN] Could not locate anchors for sections (check docx_anchor.value):")
        for m in missing:
            print(" -", m)

    # Build a sorted list of (start_idx, section)
    sec_with_idx = []
    for sec in sections:
        if sec["id"] in anchors:
            sec_with_idx.append((anchors[sec["id"]], sec))
    sec_with_idx.sort(key=lambda x: x[0])

    # Pre-extract table texts in order for later consumption
    tables_in_order: List[List[str]] = []
    for blk in iter_block_items(doc):
        if isinstance(blk, Table):
            cell_texts = []
            for row in blk.rows:
                for cell in row.cells:
                    txt = clean_text("\n".join([p.text for p in cell.paragraphs if (p.text or "").strip()]))
                    if txt:
                        cell_texts.append(txt)
            tables_in_order.append(cell_texts)

    tbl_cursor = 0

    def consume_table_texts() -> List[str]:
        nonlocal tbl_cursor
        if tbl_cursor >= len(tables_in_order):
            return []
        t = tables_in_order[tbl_cursor]
        tbl_cursor += 1
        return t

    # Extract text runs per section, including tables in range
    locale_texts = repo_root / "digcomp3-l10n" / "locale" / "texts"
    en_path = locale_texts / "en.csv"
    nl_path = locale_texts / "nl.csv"

    new_rows: List[Dict[str, str]] = []
    key_counter = 0

    for i, (start, sec) in enumerate(sec_with_idx):
        end = sec_with_idx[i+1][0] if i+1 < len(sec_with_idx) else len(blocks)
        sec_id = sec["id"]
        prefix = (sec.get("import") or {}).get("key_prefix") or f"doc.{sec_id}"

        # Walk blocks from start+1 to end-1 (exclude heading itself)
        for j in range(start+1, end):
            kind, style, txt = blocks[j]
            if kind == "p":
                if not txt:
                    continue
                key_counter += 1
                loc = f"{prefix}.u{key_counter:04d}"
                new_rows.append({
                    "location": loc,
                    "source": txt,
                    "target": "",
                    "context": f"{sec_id} | paragraph",
                })
            elif kind == "tbl":
                # consume next table texts
                cell_texts = consume_table_texts()
                for ctxt in cell_texts:
                    if not ctxt:
                        continue
                    key_counter += 1
                    loc = f"{prefix}.t{key_counter:04d}"
                    new_rows.append({
                        "location": loc,
                        "source": ctxt,
                        "target": "",
                        "context": f"{sec_id} | table-cell",
                    })

    # write en.csv and merge into nl.csv, preserving NL
    upsert_locale_csv(en_path, nl_path, new_rows)
    print("[OK] Step3 complete.")


# ----------------------------
# CLI
# ----------------------------

def main() -> None:
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="step", required=True)

    p1 = sub.add_parser("step1", help="Extract from XLSX into repo locale CSVs")
    p1.add_argument("--xlsx", required=True)
    p1.add_argument("--repo-root", default=".")

    p2 = sub.add_parser("step2", help="Check consistency with JSON-LD")
    p2.add_argument("--jsonld", required=True)
    p2.add_argument("--repo-root", default=".")

    p3 = sub.add_parser("step3", help="Extract DOCX text blocks based on manifest")
    p3.add_argument("--docx", required=True)
    p3.add_argument("--repo-root", default=".")
    p3.add_argument("--manifest", default=None)

    args = ap.parse_args()
    repo_root = Path(args.repo_root).resolve()

    if args.step == "step1":
        step1_extract_xlsx(Path(args.xlsx), repo_root)
    elif args.step == "step2":
        step2_check_jsonld(Path(args.jsonld), repo_root)
    elif args.step == "step3":
        mp = Path(args.manifest).resolve() if args.manifest else None
        step3_extract_texts(Path(args.docx), repo_root, mp)


if __name__ == "__main__":
    main()
