#!/usr/bin/env python3
"""
build_digcomp_nl_artifacts_v3.py

Build Dutch deliverables from the translated CSVs in ./digcomp3-l10n/locale:

  a) XLSX (derived from the original DigComp XLSX)
  b) JSON-LD (derived from the original DigComp JSON-LD)
  c) DOCX (rebuilt using manifest.json + translated CSVs + original XLSX for structure)

Key improvements vs v2
- XLSX: add NL columns for *all* relevant textual columns across sheets (incl. repeated area/competence names)
- JSON-LD: translate additional text fields (proficiency levels fields + proficiency level names in outcome/statement nodes)
- DOCX: include richer, complete tables for framework, levels, statements, and learning outcomes (using original XLSX rows)
- Conservative fallbacks: if a Dutch translation is missing, fall back to English source (no data loss)

Usage examples:
  python build_digcomp_nl_artifacts_v3.py --build xlsx --src-xlsx "DigComp 3.0 Data Supplement 24 Nov 2025.xlsx"
  python build_digcomp_nl_artifacts_v3.py --build jsonld --src-jsonld "DigComp 3.0 Data Supplement 24 Nov 2025.jsonld"
  python build_digcomp_nl_artifacts_v3.py --build docx --src-xlsx "DigComp 3.0 Data Supplement 24 Nov 2025.xlsx"

Assumptions
- Run this script from the *working folder* where your original sources are.
- Translation repository lives in ./digcomp3-l10n/ and contains ./locale/<component>/{en.csv,nl.csv}
- Output files are written to ./nl/ by default.

CSV format
- Expected columns: location, source, target, context
"""

from __future__ import annotations

import argparse
import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional

from openpyxl import load_workbook
from docx import Document


# ----------------------------
# CSV translation loading
# ----------------------------

def load_component_csv(path: Path) -> Dict[str, Dict[str, str]]:
    """Return mapping: location -> row dict."""
    out: Dict[str, Dict[str, str]] = {}
    if not path.exists():
        return out
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            loc = (row.get("location") or "").strip()
            if not loc:
                continue
            # normalise header variants
            out[loc] = {
                "source": row.get("source", "") or "",
                "target": row.get("target", "") or "",
                "context": row.get("context", "") or "",
            }
    return out


def load_translations(repo_root: Path) -> Dict[str, Dict[str, Dict[str, str]]]:
    """Load all locale components from ./digcomp3-l10n/locale/*/{en.csv,nl.csv}."""
    base = repo_root / "digcomp3-l10n" / "locale"
    tx: Dict[str, Dict[str, Dict[str, str]]] = {}
    if not base.exists():
        raise FileNotFoundError(f"Missing locale folder: {base}")

    for comp_dir in sorted([p for p in base.iterdir() if p.is_dir()]):
        comp = comp_dir.name
        en = load_component_csv(comp_dir / "en.csv")
        nl = load_component_csv(comp_dir / "nl.csv")
        merged: Dict[str, Dict[str, str]] = {}
        # union keys
        keys = set(en.keys()) | set(nl.keys())
        for k in keys:
            merged[k] = {
                "source": (en.get(k) or nl.get(k) or {}).get("source", "") if False else (en.get(k) or {}).get("source","") or (nl.get(k) or {}).get("source",""),
                "target": (nl.get(k) or {}).get("target", ""),
                "context": (en.get(k) or nl.get(k) or {}).get("context",""),
            }
            # if nl has no 'source', keep en source
            if not merged[k]["source"]:
                merged[k]["source"] = (en.get(k) or {}).get("source","") or (nl.get(k) or {}).get("source","")
            if not merged[k]["context"]:
                merged[k]["context"] = (en.get(k) or {}).get("context","") or (nl.get(k) or {}).get("context","")
        tx[comp] = merged
    return tx


def tr(comp: Dict[str, Dict[str, str]], key: str) -> str:
    """Dutch translation with English fallback."""
    row = comp.get(key)
    if not row:
        return ""
    tgt = (row.get("target") or "").strip()
    if tgt:
        return tgt
    return (row.get("source") or "").strip()


# ----------------------------
# Helpers
# ----------------------------

_num_re = re.compile(r"^\d+(\.\d+)?$")


def norm_num(x) -> str:
    if x is None:
        return ""
    if isinstance(x, int):
        return str(x)
    if isinstance(x, float):
        if x.is_integer():
            return str(int(x))
        return f"{x:.10f}".rstrip("0").rstrip(".")
    s = str(x).strip()
    # strip trailing .0
    if re.match(r"^\d+\.0$", s):
        return s[:-2]
    return s


def slugify_term(term: str) -> str:
    s = term.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s


def build_levels_name_map(levels: Dict[str, Dict[str, str]]) -> Dict[str, str]:
    """
    Map English four-level name -> Dutch translation.

    We use digcomp.level.<n>.four_level_name entries (n=1..8).
    """
    m: Dict[str, str] = {}
    for k, row in levels.items():
        if k.endswith(".four_level_name"):
            en = (row.get("source") or "").strip()
            nl = (row.get("target") or "").strip()
            if en and nl:
                m[en.lower()] = nl
    return m


# ----------------------------
# Build XLSX
# ----------------------------

def ensure_headers(ws, desired: List[str]) -> Dict[str, int]:
    """Ensure headers exist; return mapping header->col."""
    headers = [c.value for c in ws[1]]
    colmap = {str(h).strip(): i+1 for i, h in enumerate(headers) if h is not None}
    col = ws.max_column + 1
    for h in desired:
        if h not in colmap:
            ws.cell(row=1, column=col, value=h)
            colmap[h] = col
            col += 1
    return colmap


def build_xlsx(src_xlsx: Path, out_xlsx: Path, tx: Dict[str, Dict[str, Dict[str, str]]]) -> None:
    wb = load_workbook(src_xlsx, data_only=False)

    core = tx.get("core-framework", {})
    levels = tx.get("levels", {})
    stmts = tx.get("statements", {})
    outs = tx.get("outcomes", {})
    gloss = tx.get("glossary", {})

    # -------- Sheet 1: competence areas & competences
    ws = wb["1 Competence Areas&Competences"]
    colmap = ensure_headers(ws, [
        "Competence area name (nl)",
        "Competence area descriptor (nl)",
        "Competence name (nl)",
        "Competence descriptor (nl)",
    ])
    for r in range(2, ws.max_row + 1):
        area_no = norm_num(ws.cell(r, 1).value)
        comp_no = norm_num(ws.cell(r, 4).value)
        ws.cell(r, colmap["Competence area name (nl)"], tr(core, f"digcomp.area.{area_no}.label"))
        ws.cell(r, colmap["Competence area descriptor (nl)"], tr(core, f"digcomp.area.{area_no}.description"))
        ws.cell(r, colmap["Competence name (nl)"], tr(core, f"digcomp.competence.{comp_no}.label"))
        ws.cell(r, colmap["Competence descriptor (nl)"], tr(core, f"digcomp.competence.{comp_no}.description"))

    # -------- Sheet 2: proficiency levels
    ws = wb["2 Proficiency Levels"]
    colmap = ensure_headers(ws, [
        "Proficiency level name (nl)",
        "Four level description (nl)",
        "Purpose (nl)",
        "Eight level description (nl)",
    ])
    # Original columns: A name, B four desc, C purpose, D eight map, E six map, F eight desc
    for r in range(2, ws.max_row + 1):
        eight_map = ws.cell(r, 4).value
        if eight_map is None:
            continue
        eight_s = norm_num(eight_map)
        ws.cell(r, colmap["Proficiency level name (nl)"], tr(levels, f"digcomp.level.{eight_s}.four_level_name"))
        ws.cell(r, colmap["Four level description (nl)"], tr(levels, f"digcomp.level.{eight_s}.four_level_description"))
        ws.cell(r, colmap["Purpose (nl)"], tr(levels, f"digcomp.level.{eight_s}.applies_to"))
        ws.cell(r, colmap["Eight level description (nl)"], tr(levels, f"digcomp.level.{eight_s}.eight_level_description"))

    # Prepare mapping for repeated proficiency labels (Basic/Intermediate/etc.)
    level_name_map = build_levels_name_map(levels)

    # -------- Sheet 3: competence statements
    ws = wb["3 Competence Statements"]
    colmap = ensure_headers(ws, [
        "Competence area name (nl)",
        "Competence area descriptor (nl)",
        "Competence name (nl)",
        "Competence descriptor (nl)",
        "Proficiency level name (nl)",
        "Competence statement (nl)",
    ])
    for r in range(2, ws.max_row + 1):
        area_no = norm_num(ws.cell(r, 1).value)
        comp_no = norm_num(ws.cell(r, 4).value)
        sid = (ws.cell(r, 7).value or "")
        sid = str(sid).strip()
        prof = (ws.cell(r, 9).value or "")
        prof_s = str(prof).strip()
        ws.cell(r, colmap["Competence area name (nl)"], tr(core, f"digcomp.area.{area_no}.label"))
        ws.cell(r, colmap["Competence area descriptor (nl)"], tr(core, f"digcomp.area.{area_no}.description"))
        ws.cell(r, colmap["Competence name (nl)"], tr(core, f"digcomp.competence.{comp_no}.label"))
        ws.cell(r, colmap["Competence descriptor (nl)"], tr(core, f"digcomp.competence.{comp_no}.description"))
        if prof_s:
            ws.cell(r, colmap["Proficiency level name (nl)"], level_name_map.get(prof_s.lower(), ""))
        if sid:
            ws.cell(r, colmap["Competence statement (nl)"], tr(stmts, f"digcomp.statement.{sid}"))

    # -------- Sheet 4: learning outcomes
    ws = wb["4 Learning Outcomes"]
    colmap = ensure_headers(ws, [
        "Competence area name (nl)",
        "Competence name (nl)",
        "Proficiency level (nl)",
        "Learning outcome (nl)",
    ])
    for r in range(2, ws.max_row + 1):
        area_no = norm_num(ws.cell(r, 1).value)
        comp_no = norm_num(ws.cell(r, 3).value)
        oid = (ws.cell(r, 5).value or "")
        oid = str(oid).strip()
        prof = (ws.cell(r, 7).value or "")
        prof_s = str(prof).strip()
        ws.cell(r, colmap["Competence area name (nl)"], tr(core, f"digcomp.area.{area_no}.label"))
        ws.cell(r, colmap["Competence name (nl)"], tr(core, f"digcomp.competence.{comp_no}.label"))
        if prof_s:
            ws.cell(r, colmap["Proficiency level (nl)"], level_name_map.get(prof_s.lower(), ""))
        if oid:
            ws.cell(r, colmap["Learning outcome (nl)"], tr(outs, f"digcomp.outcome.{oid}"))

    # -------- Sheet 5: glossary
    ws = wb["5 Glossary"]
    colmap = ensure_headers(ws, ["Term (nl)", "Explanation (nl)"])
    for r in range(2, ws.max_row + 1):
        term = ws.cell(r, 1).value
        term = str(term).strip() if term is not None else ""
        if not term:
            continue
        slug = slugify_term(term)
        ws.cell(r, colmap["Term (nl)"], tr(gloss, f"digcomp.glossary.{slug}.label"))
        ws.cell(r, colmap["Explanation (nl)"], tr(gloss, f"digcomp.glossary.{slug}.definition"))

    out_xlsx.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_xlsx)
    print(f"[OK] Wrote XLSX: {out_xlsx}")


# ----------------------------
# Build JSON-LD
# ----------------------------

def load_jsonld(path: Path) -> Dict:
    return json.loads(path.read_text(encoding="utf-8"))


def build_jsonld(src_jsonld: Path, out_jsonld: Path, tx: Dict[str, Dict[str, Dict[str, str]]]) -> None:
    data = load_jsonld(src_jsonld)

    core = tx.get("core-framework", {})
    levels = tx.get("levels", {})
    stmts = tx.get("statements", {})
    outs = tx.get("outcomes", {})

    # Build mapping ProficiencyLevel/<Basic|Intermediate|...> -> NL
    level_name_map = build_levels_name_map(levels)

    def prof_uri_to_nl(uri: str) -> str:
        # Examples: ProficiencyLevel/Basic, ProficiencyLevel/Basic_1
        if not uri:
            return ""
        s = uri.split("/", 1)[-1]
        s = s.split("_", 1)[0]
        return level_name_map.get(s.lower(), "")

    graph = data.get("@graph", [])
    for node in graph:
        t = node.get("@type")
        _id = node.get("@id", "")
        if t == "CompetenceArea":
            num = _id.split("/", 1)[-1]
            node["name_nl"] = tr(core, f"digcomp.area.{num}.label") or node.get("name", "")
            node["description_nl"] = tr(core, f"digcomp.area.{num}.description") or node.get("description", "")
        elif t == "Competence":
            cid = _id.split("/", 1)[-1]
            node["name_nl"] = tr(core, f"digcomp.competence.{cid}.label") or node.get("name", "")
            node["description_nl"] = tr(core, f"digcomp.competence.{cid}.description") or node.get("description", "")
        elif t == "CompetenceStatement":
            sid = _id.split("/", 1)[-1]
            node["description_nl"] = tr(stmts, f"digcomp.statement.{sid}") or node.get("description", "")
            # translate proficiency label, if present
            p = node.get("four_levels_proficiency_name")
            if isinstance(p, str) and p:
                node["four_levels_proficiency_name_nl"] = prof_uri_to_nl(p)
        elif t == "LearningOutcome":
            oid = _id.split("/", 1)[-1]
            node["description_nl"] = tr(outs, f"digcomp.outcome.{oid}") or node.get("description", "")
            p = node.get("four_levels_proficiency_name")
            if isinstance(p, str) and p:
                node["four_levels_proficiency_name_nl"] = prof_uri_to_nl(p)
        elif t == "ProficiencyLevel":
            eight = node.get("eight_levels_mapping")
            eight_s = str(int(eight)) if isinstance(eight, (int,)) or (isinstance(eight, float) and float(eight).is_integer()) else str(eight).strip()
            # translate level fields
            node["four_levels_name_nl"] = tr(levels, f"digcomp.level.{eight_s}.four_level_name") or node.get("four_levels_name", "")
            node["four_levels_description_nl"] = tr(levels, f"digcomp.level.{eight_s}.four_level_description") or node.get("four_levels_description", "")
            node["eight_levels_description_nl"] = tr(levels, f"digcomp.level.{eight_s}.eight_level_description") or node.get("eight_levels_description", "")
            node["applies_to_nl"] = tr(levels, f"digcomp.level.{eight_s}.applies_to") or node.get("applies_to", "")

    out_jsonld.parent.mkdir(parents=True, exist_ok=True)
    out_jsonld.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] Wrote JSON-LD: {out_jsonld}")


# ----------------------------
# Build DOCX
# ----------------------------

def load_manifest(repo_root: Path) -> Dict:
    mpath = repo_root / "manifest.json"
    if not mpath.exists():
        raise FileNotFoundError(f"Missing manifest.json at {mpath}")
    return json.loads(mpath.read_text(encoding="utf-8"))


def emit_kv_paragraph(doc: Document, val: str, ctx: str = "") -> None:
    if ctx:
        p = doc.add_paragraph()
        run = p.add_run(ctx.strip())
        run.italic = True
    doc.add_paragraph(val)


def build_docx(out_docx: Path, repo_root: Path, src_xlsx: Path, tx: Dict[str, Dict[str, Dict[str, str]]]) -> None:
    manifest = load_manifest(repo_root)
    sections = manifest.get("sections", [])
    texts = tx.get("texts", {})

    core = tx.get("core-framework", {})
    levels = tx.get("levels", {})
    stmts = tx.get("statements", {})
    outs = tx.get("outcomes", {})
    gloss = tx.get("glossary", {})
    acr = tx.get("acronyms", {})

    wb = load_workbook(src_xlsx, data_only=True)
    ws_stmt = wb["3 Competence Statements"]
    ws_out = wb["4 Learning Outcomes"]
    ws_lvl = wb["2 Proficiency Levels"]

    # Map proficiency name (English) -> Dutch
    level_name_map = build_levels_name_map(levels)

    doc = Document()
    doc.add_heading("DigComp 3.0 – Nederlandse vertaling (werkversie)", level=0)

    def emit_text_prefix(prefix: str):
        # doc.<section>.<something>
        items = [(k, texts[k]) for k in texts.keys() if k.startswith(prefix + ".")]
        items.sort(key=lambda x: x[0])
        for k, row in items:
            val = (row.get("target") or "").strip() or (row.get("source") or "").strip()
            if not val:
                continue
            ctx = (row.get("context") or "").strip()
            # If context present, keep it as italic line (helps reassembly/debugging)
            if ctx:
                p = doc.add_paragraph()
                run = p.add_run(ctx)
                run.italic = True
            doc.add_paragraph(val)

    def emit_acronyms_table():
        # Try to pair "label" and "definition"; if only one, just output
        pairs: Dict[str, Dict[str, str]] = {}
        for k, row in acr.items():
            if not k.startswith("digcomp.acronym."):
                continue
            base = k.rsplit(".", 1)[0]
            pairs.setdefault(base, {})
            if k.endswith(".label"):
                pairs[base]["label"] = tr(acr, k)
            elif k.endswith(".definition"):
                pairs[base]["definition"] = tr(acr, k)
            else:
                # fallback: store as value
                pairs[base].setdefault("other", []).append(tr(acr, k))
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Acroniem"
        table.rows[0].cells[1].text = "Betekenis"
        for base in sorted(pairs.keys()):
            label = pairs[base].get("label") or base.split(".")[-1]
            defi = pairs[base].get("definition") or ""
            if not label and not defi:
                continue
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = defi

    def emit_glossary_table():
        pairs: Dict[str, Dict[str, str]] = {}
        for k in gloss.keys():
            if not k.startswith("digcomp.glossary."):
                continue
            base = k.rsplit(".", 1)[0]
            pairs.setdefault(base, {})
            if k.endswith(".label"):
                pairs[base]["label"] = tr(gloss, k)
            elif k.endswith(".definition"):
                pairs[base]["definition"] = tr(gloss, k)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Term"
        table.rows[0].cells[1].text = "Definitie"
        for base in sorted(pairs.keys()):
            term = pairs[base].get("label", "")
            defi = pairs[base].get("definition", "")
            if not term and not defi:
                continue
            cells = table.add_row().cells
            cells[0].text = term
            cells[1].text = defi

    def emit_core_framework_tables():
        # Areas
        areas = []
        for k in core.keys():
            if k.startswith("digcomp.area.") and k.endswith(".label"):
                n = k.split(".")[2]
                areas.append((n, tr(core, k), tr(core, f"digcomp.area.{n}.description")))
        areas.sort(key=lambda x: float(x[0]) if _num_re.match(x[0]) else x[0])
        doc.add_heading("Competence areas", level=2)
        t1 = doc.add_table(rows=1, cols=3)
        t1.rows[0].cells[0].text = "Nr"
        t1.rows[0].cells[1].text = "Naam"
        t1.rows[0].cells[2].text = "Beschrijving"
        for n, label, desc in areas:
            cells = t1.add_row().cells
            cells[0].text = n
            cells[1].text = label
            cells[2].text = desc

        # Competences
        comps = []
        for k in core.keys():
            if k.startswith("digcomp.competence.") and k.endswith(".label"):
                n = k.split(".")[2]
                comps.append((n, tr(core, k), tr(core, f"digcomp.competence.{n}.description")))
        def sort_key(cid: str):
            parts = cid.split(".")
            out = []
            for p in parts:
                try:
                    out.append(float(p))
                except Exception:
                    out.append(p)
            return out
        comps.sort(key=lambda x: sort_key(x[0]))
        doc.add_heading("Competences", level=2)
        t2 = doc.add_table(rows=1, cols=3)
        t2.rows[0].cells[0].text = "Nr"
        t2.rows[0].cells[1].text = "Naam"
        t2.rows[0].cells[2].text = "Beschrijving"
        for n, label, desc in comps:
            cells = t2.add_row().cells
            cells[0].text = n
            cells[1].text = label
            cells[2].text = desc

    def emit_levels_table():
        doc.add_heading("Proficiency levels", level=2)
        t = doc.add_table(rows=1, cols=6)
        t.rows[0].cells[0].text = "8-level"
        t.rows[0].cells[1].text = "4-level naam"
        t.rows[0].cells[2].text = "4-level beschrijving"
        t.rows[0].cells[3].text = "8-level beschrijving"
        t.rows[0].cells[4].text = "Doel (applies to)"
        t.rows[0].cells[5].text = "6-level"
        for r in range(2, ws_lvl.max_row + 1):
            eight = ws_lvl.cell(r, 4).value
            six = ws_lvl.cell(r, 5).value
            if eight is None:
                continue
            eight_s = norm_num(eight)
            cells = t.add_row().cells
            cells[0].text = eight_s
            cells[1].text = tr(levels, f"digcomp.level.{eight_s}.four_level_name") or (ws_lvl.cell(r,1).value or "").strip()
            cells[2].text = tr(levels, f"digcomp.level.{eight_s}.four_level_description") or (ws_lvl.cell(r,2).value or "").strip()
            cells[3].text = tr(levels, f"digcomp.level.{eight_s}.eight_level_description") or (ws_lvl.cell(r,6).value or "").strip()
            cells[4].text = tr(levels, f"digcomp.level.{eight_s}.applies_to") or (ws_lvl.cell(r,3).value or "").strip()
            cells[5].text = str(six or "").strip()

    def emit_statements_table():
        doc.add_heading("Competence statements", level=2)
        t = doc.add_table(rows=1, cols=7)
        hdr = t.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Area"
        hdr[2].text = "Competence"
        hdr[3].text = "Level"
        hdr[4].text = "AI"
        hdr[5].text = "Statement (EN)"
        hdr[6].text = "Statement (NL)"
        for r in range(2, ws_stmt.max_row + 1):
            sid = str(ws_stmt.cell(r, 7).value or "").strip()
            if not sid:
                continue
            area_no = norm_num(ws_stmt.cell(r, 1).value)
            comp_no = norm_num(ws_stmt.cell(r, 4).value)
            level_en = str(ws_stmt.cell(r, 9).value or "").strip()
            ai = str(ws_stmt.cell(r, 10).value or "").strip()
            en_txt = str(ws_stmt.cell(r, 8).value or "").strip()
            nl_txt = tr(stmts, f"digcomp.statement.{sid}") or en_txt
            cells = t.add_row().cells
            cells[0].text = sid
            cells[1].text = tr(core, f"digcomp.area.{area_no}.label") or str(ws_stmt.cell(r, 2).value or "")
            cells[2].text = tr(core, f"digcomp.competence.{comp_no}.label") or str(ws_stmt.cell(r, 5).value or "")
            cells[3].text = level_name_map.get(level_en.lower(), "") if level_en else ""
            cells[4].text = ai
            cells[5].text = en_txt
            cells[6].text = nl_txt

    def emit_outcomes_table():
        doc.add_heading("Learning outcomes", level=2)
        t = doc.add_table(rows=1, cols=8)
        hdr = t.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Area"
        hdr[2].text = "Competence"
        hdr[3].text = "Level"
        hdr[4].text = "Type"
        hdr[5].text = "AI"
        hdr[6].text = "Outcome (EN)"
        hdr[7].text = "Outcome (NL)"
        for r in range(2, ws_out.max_row + 1):
            oid = str(ws_out.cell(r, 5).value or "").strip()
            if not oid:
                continue
            area_no = norm_num(ws_out.cell(r, 1).value)
            comp_no = norm_num(ws_out.cell(r, 3).value)
            level_en = str(ws_out.cell(r, 7).value or "").strip()
            typ = str(ws_out.cell(r, 8).value or "").strip()
            ai = str(ws_out.cell(r, 9).value or "").strip()
            en_txt = str(ws_out.cell(r, 6).value or "").strip()
            nl_txt = tr(outs, f"digcomp.outcome.{oid}") or en_txt
            cells = t.add_row().cells
            cells[0].text = oid
            cells[1].text = tr(core, f"digcomp.area.{area_no}.label") or str(ws_out.cell(r, 2).value or "")
            cells[2].text = tr(core, f"digcomp.competence.{comp_no}.label") or str(ws_out.cell(r, 4).value or "")
            cells[3].text = level_name_map.get(level_en.lower(), "") if level_en else ""
            cells[4].text = typ
            cells[5].text = ai
            cells[6].text = en_txt
            cells[7].text = nl_txt

    # Drive through manifest
    for sec in sections:
        sec_id = sec.get("id", "")
        title = sec.get("title") or sec_id
        action = sec.get("action", "import")
        doc.add_heading(title, level=1)

        if action in ("import", "import_export"):
            imp = sec.get("import") or {}
            prefix = imp.get("key_prefix")
            if prefix:
                emit_text_prefix(prefix)

        if action in ("export", "import_export"):
            for ex in (sec.get("export") or []):
                comp = ex.get("component")
                if comp == "core-framework":
                    emit_core_framework_tables()
                elif comp == "levels":
                    emit_levels_table()
                elif comp == "statements":
                    emit_statements_table()
                elif comp == "outcomes":
                    emit_outcomes_table()
                elif comp == "glossary":
                    emit_glossary_table()
                elif comp == "acronyms":
                    emit_acronyms_table()
                else:
                    # fallback: dump key/value list
                    txc = tx.get(comp, {})
                    doc.add_heading(f"Export: {comp}", level=2)
                    keys = sorted(txc.keys())
                    for k in keys:
                        val = tr(txc, k)
                        if val:
                            doc.add_paragraph(f"{k}: {val}")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    print(f"[OK] Wrote DOCX: {out_docx}")


# ----------------------------
# CLI
# ----------------------------

def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--repo-root", default=".", help="Folder containing digcomp3-l10n/ and manifest.json")
    ap.add_argument("--out-dir", default="nl", help="Output folder (default: ./nl)")
    ap.add_argument("--build", choices=["xlsx", "jsonld", "docx"], required=True)
    ap.add_argument("--src-xlsx", default="DigComp 3.0 Data Supplement 24 Nov 2025.xlsx")
    ap.add_argument("--src-jsonld", default="DigComp 3.0 Data Supplement 24 Nov 2025.jsonld")
    args = ap.parse_args()

    repo_root = Path(args.repo_root).resolve()
    out_dir = (repo_root / args.out_dir).resolve()
    tx = load_translations(repo_root)

    if args.build == "xlsx":
        build_xlsx(Path(args.src_xlsx), out_dir / "DigComp_3.0_Data_Supplement_nl.xlsx", tx)
    elif args.build == "jsonld":
        build_jsonld(Path(args.src_jsonld), out_dir / "DigComp_3.0_Data_Supplement_nl.jsonld", tx)
    elif args.build == "docx":
        build_docx(out_dir / "DigComp_3.0_nl_draft.docx", repo_root, Path(args.src_xlsx), tx)


if __name__ == "__main__":
    main()
