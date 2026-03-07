import argparse
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Colors for DigComp competence areas
COLOR_MAP = {
    "1. ": "FFD966",  # Gold/Yellow
    "2. ": "8EA9DB",  # Blue
    "3. ": "F4B084",  # Orange
    "4. ": "A9D08E",  # Green
    "5. ": "F4B084",  # Red/Coral (Wait, 5 was E26B67 in the python script, let me use E26B67)
    "Standard": "1F4E78" # Standard DigComp Blue
}

def set_cell_background(cell, fill_color):
    """Set the background color of a cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def style_tables(docx_path):
    print(f"Post-processing tables in {docx_path}...")
    doc = docx.Document(docx_path)
    
    # Define area colors
    colors = {
        "1. ": "FFD966",
        "2. ": "8EA9DB",
        "3. ": "F4B084",
        "4. ": "A9D08E",
        "5. ": "E26B67"
    }

    for table in doc.tables:
        for row in table.rows:
            # Infer area by looking at the second column's ID (e.g. "1.1 Browsen")
            # Or if it's the specific outcomes table, maybe the first column "LO1.1.X"?
            area_to_color = None
            if len(row.cells) > 1:
                col1_text = row.cells[0].text.strip()
                col2_text = row.cells[1].text.strip()
                
                # Check directly in the row's cells for our Area keywords:
                combined_text = (col1_text + " " + col2_text).replace('\n', ' ')
                
                if "1." in combined_text and ("1. INFORMATIE" in combined_text or col2_text.startswith("1.1 ") or col1_text.startswith("LO1.")):
                    area_to_color = "FFD966"
                elif "2." in combined_text and ("2. COMMUNICATIE" in combined_text or col2_text.startswith("2.1 ") or col1_text.startswith("LO2.")):
                    area_to_color = "8EA9DB"
                elif "3." in combined_text and ("3. CONTENTCREATIE" in combined_text or col2_text.startswith("3.1 ") or col1_text.startswith("LO3.")):
                    area_to_color = "F4B084"
                elif "4." in combined_text and ("4. VEILIGHEID" in combined_text or col2_text.startswith("4.1 ") or col1_text.startswith("LO4.")):
                    area_to_color = "A9D08E"
                elif "5." in combined_text and ("5. PROBLEEMOPLOSSING" in combined_text or col2_text.startswith("5.1 ") or col1_text.startswith("LO5.")):
                    area_to_color = "E26B67"

            for cell in row.cells:
                # 1. Left align all text in the cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                if area_to_color:
                    set_cell_background(cell, area_to_color)
                    
    # Save the modified document
    doc.save(docx_path)
    print(f"Successfully styled tables in {docx_path}.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("file", help="Path to the docx file to style")
    args = parser.parse_args()
    style_tables(args.file)
