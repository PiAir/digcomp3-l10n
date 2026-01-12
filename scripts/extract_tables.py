import pandas as pd
import json
import argparse
import os
import re
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# --- CONFIGURATIE EN KLEUREN ---

COLOR_MAP = {
    "CompetenceArea/1": ["FFF2CC", "FFE699", "FFD966", "FFC000"], # Goud/Geel
    "CompetenceArea/2": ["D9E1F2", "B4C6E7", "8EA9DB", "2F5597"], # Blauw
    "CompetenceArea/3": ["FBE5D6", "F8CBAD", "F4B084", "ED7D31"], # Oranje
    "CompetenceArea/4": ["E2EFDA", "C6E0B4", "A9D08E", "70AD47"], # Groen
    "CompetenceArea/5": ["FCE4D6", "F9CB9C", "F4B084", "E26B67"], # Rood/Coral
    "Standard": "1F4E78" # Standaard DigComp Blauw
}

AREA_COLORS = {k: v[3] for k, v in COLOR_MAP.items() if k != "Standard"}

# --- OPMAAK HULPFUNCTIES ---

def set_cell_background(cell, fill):
    """Zet achtergrondkleur van een cel (Hex string zonder #)."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_repeat_header(row):
    """Zorgt dat de rij herhaald wordt bovenaan elke pagina."""
    trPr = row._tr.get_or_add_trPr()
    tHeader = parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />')
    trPr.append(tHeader)

def set_no_split(row):
    """Voorkomt dat een rij wordt gesplitst over twee pagina's."""
    trPr = row._tr.get_or_add_trPr()
    keepNext = parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="true"/>')
    trPr.append(keepNext)

def set_table_width_100(table):
    """Forceert de tabel op 100% breedte (Aanpassen aan venster)."""
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000') 
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def style_text(paragraph, text, bold=False, italic=False, size=9, color="000000"):
    """Past Arial stijl toe op tekst binnen een paragraaf."""
    run = paragraph.add_run(str(text))
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def apply_pdf_borders(table, sz='2'):
    """Past subtiele horizontale lijnen toe zoals in de PDF."""
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border in ['left', 'right', 'insideV']:
        b = OxmlElement(f'w:{border}'); b.set(qn('w:val'), 'nil'); tblBorders.append(b)
    for border in ['top', 'bottom', 'insideH']:
        b = OxmlElement(f'w:{border}'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz); b.set(qn('w:color'), 'D9D9D9'); tblBorders.append(b)
    tblPr.append(tblBorders)

def format_ai_label(label_raw, lang):
    """Zet de ruwe AI label tekst om naar voluit geschreven varianten."""
    if not label_raw or "not Implicit" in label_raw or label_raw == "-":
        return "-"
    if lang == "nl":
        return label_raw.replace('AI-Implicit', 'AI-Impliciet').replace('AI-Explicit', 'AI-Expliciet')
    return label_raw

# --- GENERATIE FUNCTIES ---

def generate_csv_output(csv_path, type_name, lang, output_path):
    """Verwerkt Acronyms of Glossary vanuit CSV."""
    df = pd.read_csv(csv_path, encoding='utf-8-sig')
    grouped = {}
    search_key = "acronym" if type_name == "acronyms" else "glossary"
    content_idx = 2 if lang == "en" else 3
    
    pattern = re.compile(rf"{search_key}\.([^.]+)\.(label|expansion|definition|source)", re.I)
    for _, row in df.iterrows():
        key, val = str(row.iloc[0]), str(row.iloc[content_idx])
        match = pattern.search(key)
        if match:
            id_str, field = match.group(1), match.group(2).lower()
            if id_str not in grouped: grouped[id_str] = {'c1': id_str.upper(), 'c2': '', 'c3': ''}
            if field == 'label': grouped[id_str]['c1'] = val
            elif field in ['expansion', 'definition']: grouped[id_str]['c2'] = val
            elif field == 'source': grouped[id_str]['c3'] = val

    data = sorted(grouped.values(), key=lambda x: x['c1'].lower())
    doc = Document()
    titles = {"acronyms": ("LIST OF ACRONYMS", "LIJST VAN ACRONYMEN"), "glossary": ("GLOSSARY OF TERMS AND DEFINITIONS", "GLOSSARIUM VAN TERMEN EN DEFINITIES")}
    style_text(doc.add_heading('', level=1), titles[type_name][0] if lang == "en" else titles[type_name][1], size=18, color=COLOR_MAP["Standard"])
    
    table = doc.add_table(rows=0, cols=4 if type_name == "acronyms" else 3)
    set_table_width_100(table); apply_pdf_borders(table)
    
    headers = (["ACRONYM", "DESCRIPTION", "ACRONYM", "DESCRIPTION"] if type_name == "acronyms" else ["TERM", "EXPLANATION", "SOURCE"]) if lang == "en" else \
              (["ACRONIEM", "BETEKENIS", "ACRONIEM", "BETEKENIS"] if type_name == "acronyms" else ["TERM", "UITLEG", "BRON"])
    
    h_row = table.add_row().cells
    for i, h in enumerate(headers):
        style_text(h_row[i].paragraphs[0], h, bold=True, color="FFFFFF")
        set_cell_background(h_row[i], COLOR_MAP["Standard"])

    if type_name == "acronyms":
        mid = (len(data) + 1) // 2
        for i in range(mid):
            r = table.add_row().cells
            style_text(r[0].paragraphs[0], data[i]['c1'], bold=True)
            style_text(r[1].paragraphs[0], data[i]['c2'])
            if i + mid < len(data):
                style_text(r[2].paragraphs[0], data[i+mid]['c1'], bold=True)
                style_text(r[3].paragraphs[0], data[i+mid]['c2'])
    else:
        for item in data:
            r = table.add_row().cells
            style_text(r[0].paragraphs[0], item['c1'], bold=True); style_text(r[1].paragraphs[0], item['c2']); style_text(r[2].paragraphs[0], item['c3'])
    doc.save(output_path)

def generate_digcomp3(json_path, lang, output_path, images_path):
    """Genereert de gedetailleerde 3.2 competentiepagina's."""
    with open(json_path, 'r', encoding='utf-8') as f:
        graph = json.load(f)["@graph"]
    doc = Document()
    doc.sections[0].left_margin, doc.sections[0].right_margin = Cm(1.2), Cm(1.2)
    areas = sorted([i for i in graph if i["@type"] == "CompetenceArea"], key=lambda x: x["@id"])
    competences = [i for i in graph if i["@type"] == "Competence"]; levels = [i for i in graph if i["@type"] == "ProficiencyLevel"]
    statements = [i for i in graph if i["@type"] == "CompetenceStatement"]; suffix = "_nl" if lang == "nl" else ""
    for area in areas:
        area_id = area["@id"]; area_num = area_id.split("/")[-1]; tints = COLOR_MAP.get(area_id, ["FFFFFF"]*4)
        for comp in sorted([c for c in competences if c["competence_area_id"] == area_id], key=lambda x: x["@id"]):
            comp_id_short = comp["@id"].split("/")[-1]
            table = doc.add_table(rows=4, cols=3); table.autofit = False; set_table_width_100(table); apply_pdf_borders(table)
            for i, w in enumerate([Cm(5.0), Cm(3.8), Cm(9.8)]): table.columns[i].width = w
            left_cell = table.cell(0, 0).merge(table.cell(3, 0))
            img_path = os.path.join(images_path, f"DC3_{comp_id_short.replace('.', 'p')}.png")
            if os.path.exists(img_path): left_cell.paragraphs[0].add_run().add_picture(img_path, width=Cm(3.8))
            style_text(left_cell.add_paragraph(), f"{area_num}. {area[f'name{suffix}'].upper()}", bold=True, size=10, color=tints[3])
            style_text(left_cell.add_paragraph(), f"{comp_id_short} {comp[f'name{suffix}']}", bold=True, size=11)
            style_text(left_cell.add_paragraph(), comp[f"description{suffix}"], size=8.5)
            for idx, lv_key in enumerate(["Basic", "Intermediate", "Advanced", "Highly advanced"]):
                bg_color = tints[idx]; txt_color = "FFFFFF" if area_id == "CompetenceArea/2" and idx > 1 else "000000"
                lvl_cell = table.cell(idx, 1); set_cell_background(lvl_cell, bg_color)
                lv_name = next((l[f"four_levels_name{suffix}"] for l in levels if lv_key in l["@id"]), lv_key)
                lvl_text = f"At {lv_name} level, individuals" if lang == "en" else f"Op {lv_name.lower()}niveau kunnen individuen"
                style_text(lvl_cell.paragraphs[0], lvl_text, bold=True, italic=True, color=txt_color)
                stmt_cell = table.cell(idx, 2)
                relevant = [s for s in statements if s["competence_id"] == comp["@id"] and s["four_levels_proficiency_name"].startswith(f"ProficiencyLevel/{lv_key}")]
                for s in relevant:
                    p = stmt_cell.add_paragraph(); ai = format_ai_label(s.get('ai_label', ''), lang)
                    style_text(p, f"{s['@id'].split('/')[-1]}: ", bold=True); style_text(p, f"{s[f'description{suffix}']} "); style_text(p, f"[{ai}]", bold=True)
            doc.add_paragraph()
    doc.save(output_path)

def generate_table2(json_path, lang, output_path, images_path):
    """Genereert Tabel 2: Overzicht van gebieden en competenties."""
    with open(json_path, 'r', encoding='utf-8') as f:
        graph = json.load(f)["@graph"]
    doc = Document(); doc.sections[0].left_margin, doc.sections[0].right_margin = Cm(1.2), Cm(1.2)
    areas = sorted([i for i in graph if i["@type"] == "CompetenceArea"], key=lambda x: x["@id"])
    competences = [i for i in graph if i["@type"] == "Competence"]; suffix = "_nl" if lang == "nl" else ""
    table = doc.add_table(rows=1, cols=3); table.autofit = False; set_table_width_100(table)
    h_cells = table.rows[0].cells
    headers = ["AREA", "TITLE", "DESCRIPTOR"] if lang == "en" else ["GEBIED", "TITEL", "BESCHRIJVING"]
    for i, txt in enumerate(headers): style_text(h_cells[i].paragraphs[0], txt, bold=True, color="2F5597")
    for area in areas:
        area_id = area["@id"]; area_num = area_id.split("/")[-1]; color = AREA_COLORS.get(area_id, "000000")
        area_comps = sorted([c for c in competences if c["competence_area_id"] == area_id], key=lambda x: x["@id"])
        start_row = len(table.rows)
        for comp in area_comps:
            r = table.add_row().cells; style_text(r[1].paragraphs[0], f"{comp['@id'].split('/')[-1]} {comp[f'name{suffix}']}", bold=True); style_text(r[2].paragraphs[0], comp[f"description{suffix}"])
        a_cell = table.cell(start_row, 0).merge(table.cell(len(table.rows)-1, 0))
        img_path = os.path.join(images_path, f"DC3_small_c{area_num}.png")
        if os.path.exists(img_path): a_cell.paragraphs[0].add_run().add_picture(img_path, width=Cm(1.5))
        style_text(a_cell.add_paragraph(), f"{area_num}. {area[f'name{suffix}'].upper()}", bold=True, color=color); style_text(a_cell.add_paragraph(), area[f"description{suffix}"], size=8)
    apply_pdf_borders(table, sz='4'); doc.save(output_path)

def generate_outcomes(json_path, lang, output_path):
    """Genereert Learning Outcomes met herhalende koppen en 100% breedte."""
    with open(json_path, 'r', encoding='utf-8') as f:
        graph = json.load(f)["@graph"]
    doc = Document(); section = doc.sections[0]; section.left_margin, section.right_margin = Cm(1.0), Cm(1.0)
    areas = sorted([i for i in graph if i["@type"] == "CompetenceArea"], key=lambda x: x["@id"])
    competences = [i for i in graph if i["@type"] == "Competence"]
    levels = [i for i in graph if i["@type"] == "ProficiencyLevel"]
    outcomes = [i for i in graph if i["@type"] == "LearningOutcome"]; suffix = "_nl" if lang == "nl" else ""
    for area in areas:
        area_id = area["@id"]; area_label = "COMPETENCE AREA" if lang == "en" else "COMPETENTIEGEBIED"; comp_label = "Competentie" if lang == "nl" else "Competence"
        color = COLOR_MAP.get(area_id, ["FFFFFF"]*4)[3]
        for comp in sorted([c for c in competences if c["competence_area_id"] == area_id], key=lambda x: x["@id"]):
            comp_id = comp["@id"]; comp_num = comp_id.split("/")[-1]
            table = doc.add_table(rows=2, cols=5); table.autofit = False; set_table_width_100(table)
            widths = [Cm(1.5), Cm(10.0), Cm(2.5), Cm(2.5), Cm(2.5)]
            for i, w in enumerate(widths): table.columns[i].width = w
            # Rij 1: Kop (Wit op Kleur)
            header_cell = table.cell(0, 0).merge(table.cell(0, 4))
            header_text = f"{area_label} {area_id.split('/')[-1]}: {area[f'name{suffix}']} - {comp_label} {comp_num} {comp[f'name{suffix}']}"
            style_text(header_cell.paragraphs[0], header_text.upper(), bold=True, size=10, color=color)
            set_repeat_header(table.rows[0])
            # Rij 2: Labels
            labels = ["ID", "Outcome", "Level", "K/S/A", "AI"] if lang=="en" else ["ID", "Leerresultaat", "Niveau", "K/V/H", "AI-label"]
            for i, label in enumerate(labels):
                cell = table.cell(1, i); style_text(cell.paragraphs[0], label, bold=True, size=8, color="000000"); set_cell_background(cell, color)
            set_repeat_header(table.rows[1])
            for o in sorted([o for o in outcomes if o["competence_id"] == comp_id], key=lambda x: x["@id"]):
                row = table.add_row(); set_no_split(row); cells = row.cells
                style_text(cells[0].paragraphs[0], o["@id"].split("/")[-1], bold=True, size=8)
                style_text(cells[1].paragraphs[0], o[f"description{suffix}"], size=9)
                lv_name = next((l[f"four_levels_name{suffix}"] for l in levels if o["four_levels_proficiency_name"] in l["@id"]), "Level")
                style_text(cells[2].paragraphs[0], lv_name, size=8)
                type_val = o.get("type", "").replace("Knowledge", "Kennis").replace("Skill", "Vaardigheid").replace("Attitude", "Houding") if lang == "nl" else o.get("type", "")
                style_text(cells[3].paragraphs[0], type_val, size=8)
                style_text(cells[4].paragraphs[0], format_ai_label(o.get("ai_label", ""), lang), size=8)
            apply_pdf_borders(table); doc.add_paragraph()
    doc.save(output_path)

# --- MAIN ---

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('type', choices=['acronyms', 'glossary', 'digcomp3', 'table2', 'outcomes'])
    parser.add_argument('--lang', choices=['en', 'nl'], default='nl')
    parser.add_argument('--path', default=r'.\digcomp3-l10n\locale')
    parser.add_argument('--json', default=r'.\nl\DigComp_3.0_Data_Supplement_nl.jsonld')
    parser.add_argument('--images', default=r'.\digcomp3-l10n\images')
    parser.add_argument('--output', help='Output filename')
    args = parser.parse_args()

    out = args.output if args.output else f"DigComp3_{args.type}_{args.lang}.docx"
    
    if args.type == 'outcomes': generate_outcomes(args.json, args.lang, out)
    elif args.type == 'digcomp3': generate_digcomp3(args.json, args.lang, out, args.images)
    elif args.type == 'table2': generate_table2(args.json, args.lang, out, args.images)
    else:
        csv_path = os.path.join(args.path, args.type, f"{args.lang}.csv")
        generate_csv_output(csv_path, args.type, args.lang, out)
    print(f"Gereed: {out}")

if __name__ == "__main__":
    main()