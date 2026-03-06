#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_digcomp_docx_from_locale_v5.py

Fixes the two practical issues you observed:
1) Headings not translated: headings are just paragraphs; we translate them too.
2) Competence tables partly untranslated: many CSV extractions hash/translate whole-cell text.
   This builder tries WHOLE-CELL replacement first, then per-paragraph fallback.

Matching strategy (robust):
- Build hash->target index from locale/<component>/<lang>.csv using sha1(normalized(source)).
- Also uses sha1 suffix in 'location' if present (40-hex at end).
- For each paragraph: translate by full-paragraph hash; if not found and contains \\n, translate per-line.
- For each table cell: try translating the FULL CELL TEXT; if match, replace cell as one paragraph with \\n.
  Otherwise translate per-paragraph in the cell.

Usage:
python build_digcomp_docx_from_locale_v5.py ^
  --template-docx "DigComp 3.0 EDITABLE 16 Dec 2025 - bewerkt.docx" ^
  --repo-root . ^
  --out-docx ".\\nl\\DigComp_3.0_nl.docx"
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Optional, List, Tuple

from docx import Document

WS_RE = re.compile(r"[ \t\u00A0]+")

def normalize(s: str) -> str:
    if s is None:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = WS_RE.sub(" ", s)
    lines = [ln.strip(" \t\u00A0") for ln in s.split("\n")]
    while lines and lines[0] == "":
        lines.pop(0)
    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines)

def sha1_hex(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8")).hexdigest()

@dataclass
class Index:
    by_hash: Dict[str, str]
    rows: int
    rows_with_target: int

def load_index(csv_path: Path) -> Index:
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        if not reader.fieldnames:
            raise ValueError(f"No header in CSV: {csv_path}")
        headers = {h.lower(): h for h in reader.fieldnames}
        if "source" not in headers:
            raise ValueError(f"CSV missing 'source' column: {reader.fieldnames}")
        if "target" not in headers and "translation" not in headers:
            raise ValueError(f"CSV missing 'target' column: {reader.fieldnames}")
        c_source = headers["source"]
        c_target = headers.get("target") or headers.get("translation")
        c_loc = headers.get("location")

        by_hash: Dict[str, str] = {}
        rows = 0
        rows_with_target = 0

        for r in reader:
            rows += 1
            src = (r.get(c_source) or "")
            tgt = (r.get(c_target) or "")
            if not tgt.strip():
                continue
            rows_with_target += 1

            if c_loc:
                loc = (r.get(c_loc) or "").strip()
                if loc:
                    last = loc.split(".")[-1]
                    if re.fullmatch(r"[0-9a-f]{40}", last):
                        by_hash.setdefault(last, tgt)

            srcn = normalize(src)
            if srcn:
                by_hash.setdefault(sha1_hex(srcn), tgt)

    return Index(by_hash=by_hash, rows=rows, rows_with_target=rows_with_target)

def set_paragraph_text_keep_first_run(p, text: str):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for run in p.runs[1:]:
        run.text = ""

def replace_cell_as_single_paragraph(cell, text: str):
    # ensure at least one paragraph
    if not cell.paragraphs:
        p = cell.add_paragraph("")
        set_paragraph_text_keep_first_run(p, text)
        return
    set_paragraph_text_keep_first_run(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        set_paragraph_text_keep_first_run(p, "")

def translate_text(text: str, idx: Index) -> Tuple[str, bool]:
    n = normalize(text)
    if not n:
        return text, False
    h = sha1_hex(n)
    if h in idx.by_hash:
        return idx.by_hash[h], True
    # line fallback for multi-line paragraphs
    if "\n" in n:
        parts = n.split("\n")
        out = []
        changed = False
        for part in parts:
            pn = normalize(part)
            if not pn:
                out.append(part)
                continue
            hp = sha1_hex(pn)
            if hp in idx.by_hash:
                out.append(idx.by_hash[hp])
                changed = True
            else:
                out.append(part)
        if changed:
            return "\n".join(out), True
    return text, False

def full_cell_text(cell) -> str:
    parts: List[str] = []
    for p in cell.paragraphs:
        t = normalize(p.text or "")
        if t:
            parts.append(t)
    return "\n".join(parts)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template-docx", required=True)
    ap.add_argument("--repo-root", default=".")
    ap.add_argument("--out-docx", required=True)
    ap.add_argument("--component", default="texts")
    ap.add_argument("--lang", default="nl")
    ap.add_argument("--locale-dir", default=None)
    args = ap.parse_args()

    repo_root = Path(args.repo_root).resolve()
    locale_dir = Path(args.locale_dir).resolve() if args.locale_dir else (repo_root / "digcomp3-l10n" / "locale")
    csv_path = locale_dir / args.component / f"{args.lang}.csv"
    if not csv_path.exists():
        raise SystemExit(f"CSV not found: {csv_path}")

    idx = load_index(csv_path)
    doc = Document(args.template_docx)

    stats = {"paras_total":0, "paras_changed":0, "cells_total":0, "cells_changed_full":0, "cells_changed_para":0}

    # Paragraphs (includes headings)
    for p in doc.paragraphs:
        stats["paras_total"] += 1
        old = p.text or ""
        new, changed = translate_text(old, idx)
        if changed and new != old:
            set_paragraph_text_keep_first_run(p, new)
            stats["paras_changed"] += 1

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stats["cells_total"] += 1
                old_full = full_cell_text(cell)
                if not normalize(old_full):
                    continue
                new_full, changed_full = translate_text(old_full, idx)
                if changed_full and normalize(new_full) != normalize(old_full):
                    replace_cell_as_single_paragraph(cell, new_full)
                    stats["cells_changed_full"] += 1
                    continue
                # fallback per paragraph
                cell_changed = False
                for p in cell.paragraphs:
                    old = p.text or ""
                    new, changed = translate_text(old, idx)
                    if changed and new != old:
                        set_paragraph_text_keep_first_run(p, new)
                        cell_changed = True
                if cell_changed:
                    stats["cells_changed_para"] += 1

    out_path = Path(args.out_docx)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    print(f"[OK] Wrote: {out_path}")
    print(f"Index: csv_rows={idx.rows}, rows_with_target={idx.rows_with_target}, hash_index={len(idx.by_hash)}")
    print(f"Paragraphs changed: {stats['paras_changed']}/{stats['paras_total']}")
    print(f"Cells changed (full): {stats['cells_changed_full']}/{stats['cells_total']}")
    print(f"Cells changed (per-paragraph fallback): {stats['cells_changed_para']}/{stats['cells_total']}")

if __name__ == "__main__":
    main()
